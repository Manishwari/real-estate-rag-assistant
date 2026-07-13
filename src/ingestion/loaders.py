"""
Document loaders for the Real Estate Knowledge Base.

Each loader takes a file path and returns a single `Document` dict:
    {
        "doc_id": "<relative path, used as a stable id>",
        "source": "<filename>",
        "format": "pdf" | "docx" | "html" | "markdown",
        "category": "<inferred category, e.g. brochure, faq, payment_plan>",
        "text": "<full extracted plain text>",
    }

Keeping this format-agnostic and returning plain text lets the chunker
and everything downstream stay completely format-agnostic too.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class Document:
    doc_id: str
    source: str
    format: str
    category: str
    text: str
    metadata: dict = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Category inference
# ---------------------------------------------------------------------------
# The dataset encodes a lot of useful signal in the filename itself
# (e.g. "mlv_brochure.pdf", "skyline_faq.html"). We turn that into a
# human-readable category used for citations and filtering.
_CATEGORY_PATTERNS = [
    (r"brochure", "Project Brochure"),
    (r"builder_profile", "Builder Profile"),
    (r"rera_summary", "RERA Documentation"),
    (r"rera_general", "RERA Documentation"),
    (r"privacy_policy", "Privacy Policy"),
    (r"terms_conditions|terms_of_use", "Terms & Conditions"),
    (r"faq", "FAQ"),
    (r"payment_plan", "Payment Plan"),
    (r"cancellation_refund", "Cancellation & Refund Policy"),
    (r"home_loan", "Home Loan Information"),
    (r"registration_process", "Registration Process"),
    (r"possession_guidelines", "Possession Guidelines"),
    (r"customer_support", "Customer Support"),
    (r"amenities_guide", "Amenities Guide"),
    (r"location_guide", "Location Guide"),
    (r"floor_plans", "Floor Plan"),
    (r"sale_agreement", "Sale Agreement / Legal Terms"),
    (r"listing", "Property Listing"),
    (r"_home\b|_about\b", "Builder Profile"),
]


def infer_category(filename: str) -> str:
    stem = filename.lower()
    for pattern, label in _CATEGORY_PATTERNS:
        if re.search(pattern, stem):
            return label
    return "General Document"


def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Format-specific extraction
# ---------------------------------------------------------------------------
def load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _clean_text("\n\n".join(pages))


def load_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean_text("\n".join(parts))


def load_html(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return _clean_text(text)


def load_markdown(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()
    # Strip markdown formatting characters but keep the readable text
    text = re.sub(r"^#{1,6}\s*", "", raw, flags=re.MULTILINE)
    text = re.sub(r"[*_`>#-]{1,3}", "", text)
    return _clean_text(text)


_LOADERS = {
    "pdf": load_pdf,
    "docx": load_docx,
    "html": load_html,
    "markdown": load_markdown,
}

_EXT_TO_FORMAT = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".html": "html",
    ".htm": "html",
    ".md": "markdown",
    ".markdown": "markdown",
}


def load_document(path: Path, raw_root: Path) -> Document | None:
    fmt = _EXT_TO_FORMAT.get(path.suffix.lower())
    if fmt is None:
        return None
    try:
        text = _LOADERS[fmt](path)
    except Exception as exc:  # noqa: BLE001
        print(f"[loader] WARNING: failed to parse {path}: {exc}")
        return None
    if not text or len(text) < 20:
        return None
    doc_id = str(path.relative_to(raw_root))
    return Document(
        doc_id=doc_id,
        source=path.name,
        format=fmt,
        category=infer_category(path.name),
        text=text,
    )


def load_all_documents(raw_root: Path) -> list[Document]:
    """Walk data/raw/{pdf,docx,html,markdown}/ and load every file."""
    docs: list[Document] = []
    for path in sorted(raw_root.rglob("*")):
        if path.is_file():
            doc = load_document(path, raw_root)
            if doc:
                docs.append(doc)
    return docs
